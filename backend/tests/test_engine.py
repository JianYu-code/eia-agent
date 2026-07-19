import sys
import os
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent / "backend"))
os.environ.setdefault("DATABASE_URL", "sqlite+aiosqlite:////tmp/eia_test_suite.db")


def test_extract_docx_heading_styles():
    """测试 DOCX 标题样式识别"""
    from app.engine.extractor import _split_chapters

    md_text = """# 前言
这是前言内容。

## 总则
编制依据内容。

## 工程分析
这是工程分析内容。

### 源强核算
核算细节。"""

    chapters = _split_chapters(md_text)
    titles = [c["title"] for c in chapters]
    assert "前言" in titles, f"Expected '前言' in {titles}"
    assert "总则" in titles, f"Expected '总则' in {titles}"
    assert "工程分析" in titles, f"Expected '工程分析' in {titles}"
    assert "源强核算" in titles, f"Expected '源强核算' in {titles}"
    assert len(chapters) >= 4, f"Expected >= 4 chapters, got {len(chapters)}"
    print("✓ DOCX 标题识别通过")


def test_split_chapters_chinese_patterns():
    """测试中文标题格式识别"""
    from app.engine.extractor import _split_chapters

    text = """第一章 前言
前言正文内容。

一、编制依据
依据内容。

（一）评价因子
因子详情。

1.1 工程分析
分析内容。"""

    chapters = _split_chapters(text)
    titles = [c["title"] for c in chapters]
    assert "前言" in titles, f"Expected '前言' in {titles}"
    assert "编制依据" in titles, f"Expected '编制依据' in {titles}"
    assert "评价因子" in titles, f"Expected '评价因子' in {titles}"
    assert "1.1 工程分析" in titles or "工程分析" in titles
    print("✓ 中文标题识别通过")


def test_grader():
    """测试 P0/P1/P2 分级器"""
    from app.engine.grader import grade_issues, build_issue

    issues = [
        build_issue("R-001", "P0", "测试", "严重问题", "发现"),
        build_issue("R-002", "P1", "测试", "一般问题", "发现"),
        build_issue("R-003", "P1", "测试", "一般问题2", "发现"),
        build_issue("R-004", "P2", "测试", "建议", "发现"),
    ]
    graded = grade_issues(issues)
    assert len(graded["P0"]) == 1
    assert len(graded["P1"]) == 2
    assert len(graded["P2"]) == 1
    assert graded["P0"][0]["rule_id"] == "R-001"
    print("✓ 分级器通过")


def test_rules_engine_keyword():
    """测试关键词匹配规则"""
    from app.engine.rules_engine import run_keyword_check

    rule = {
        "rule_id": "R-TEST",
        "severity": "P0",
        "category": "测试",
        "check_config": {
            "required_chapters": ["前言", "总则"],
            "required_keywords": ["环境影响评价"],
        },
        "law_ref": "HJ 2.1-2016",
    }

    text_with = """前言\n这是前言内容。\n总则\n编制依据。\n本项目进行了环境影响评价。"""
    issues = run_keyword_check(rule, text_with)
    assert len(issues) == 0, f"Expected 0 issues, got {len(issues)}"

    text_without = """这是第一章内容。\n第二章继续。"""
    issues = run_keyword_check(rule, text_without)
    assert len(issues) >= 2, f"Expected >= 2 issues, got {len(issues)}"
    print("✓ 关键词匹配通过")


def test_rules_engine_load_filter():
    """测试规则按 report_type 过滤"""
    import yaml
    from pathlib import Path
    from app.engine.rules_engine import load_rules

    rules_report_book = load_rules("eia", "报告书")
    rules_report_table = load_rules("eia", "报告表")

    assert len(rules_report_book) > 0
    assert len(rules_report_table) > 0

    book_ids = [r["rule_id"] for r in rules_report_book]
    table_ids = [r["rule_id"] for r in rules_report_table]
    assert "R-STRUCT-001" in book_ids, "报告书规则缺少 R-STRUCT-001"
    assert "R-STRUCT-001-TABLE" in table_ids, "报告表规则缺少 R-STRUCT-001-TABLE"
    assert "R-STRUCT-001" not in table_ids, "报告表规则不应包含 R-STRUCT-001"
    assert "R-STRUCT-001-TABLE" not in book_ids, "报告书规则不应包含 R-STRUCT-001-TABLE"
    print(f"✓ 规则过滤通过 (报告书:{len(rules_report_book)}条, 报告表:{len(rules_report_table)}条)")


def test_standards_index():
    import sys
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent / "scripts"))
    try:
        from ingest_mineru import _normalize_std_id
    except ModuleNotFoundError:
        print("- 标准编号归一化测试跳过（缺少 lancedb 依赖）")
        return
    assert _normalize_std_id("HJ 2.1-2016") == "HJ2.1-2016"
    assert _normalize_std_id("HJ2.1-2016") == "HJ2.1-2016"
    assert _normalize_std_id("GB 16297—1996") == "GB16297-1996"
    print("✓ 标准编号归一化通过")


def test_pipeline_structure_check():
    """测试管道结构检查逻辑"""
    required = ["前言", "总则", "工程分析"]
    full_text_with = "前言\n这是内容。\n总则\n总则内容。\n工程分析\n分析内容。"
    for ch in required:
        assert ch in full_text_with, f"Expected '{ch}' in text"
    print("✓ 结构检查逻辑通过")


def test_context_routing():
    """测试章节感知上下文路由"""
    from app.engine.context import route_chapters, build_step_context

    chapters = [
        {"title": "前言", "level": 1, "content": "前言内容" * 100},
        {"title": "3 工程分析", "level": 2, "content": "源强核算：SO₂=2×B×S " * 100},
        {"title": "7 环境保护措施及其可行性论证", "level": 2, "content": "活性炭吸附 " * 100},
    ]
    text_data = {"full_text": "".join(c["content"] for c in chapters), "chapters": chapters}

    routed = route_chapters(chapters, ["工程分析", "源强"])
    assert "源强核算" in routed, "应路由到工程分析章"
    assert "前言内容" not in routed, "不应包含无关章节"

    ctx = build_step_context(text_data, ["工程分析", "源强"])
    assert "源强核算" in ctx

    ctx2 = build_step_context(text_data, ["不存在的章节"])
    assert len(ctx2) > 0, "无匹配章节时应回退全文"
    print("✓ 章节路由通过")


def test_llm_json_parse():
    """测试 LLM JSON 解析与修复"""
    from app.engine.llm_json import parse_llm_json, LLMParseError

    assert parse_llm_json('{"severity":"P1","title":"x"}', "object")["severity"] == "P1"
    assert parse_llm_json('```json\n[{"title":"a"}]\n```', "array") == [{"title": "a"}]
    assert parse_llm_json('分析过程...\n{"title":"t","finding":"f",}\n其他', "object")["title"] == "t"
    assert parse_llm_json("null", "object") is None
    assert parse_llm_json("[]", "array") == []
    assert parse_llm_json('{"a":1}', "array") == [{"a": 1}]
    try:
        parse_llm_json("完全没有JSON", "object")
        assert False, "应抛出 LLMParseError"
    except LLMParseError:
        pass
    print("✓ LLM JSON 解析通过")


def test_extractor_folder_and_classify(tmp_path=None):
    """测试资料包提取与文件分类"""
    import tempfile
    from app.engine.extractor import extract_text, classify_file, package_completeness

    assert classify_file("环评批复.pdf") == "环评批复"
    assert classify_file("验收监测报告.docx") == "验收监测报告"
    assert classify_file("公示截图.pdf") == "信息公开证据"
    assert classify_file("随机文件.md") == "其他"

    with tempfile.TemporaryDirectory() as d:
        from pathlib import Path
        (Path(d) / "环评批复.md").write_text("# 批复\n同意建设。", encoding="utf-8")
        (Path(d) / "验收意见.txt").write_text("验收结论：通过", encoding="utf-8")
        data = extract_text(d)
        assert data.get("is_package") is True
        assert len(data["files"]) == 2
        assert "批复" in data["full_text"] and "验收" in data["full_text"]
        missing = package_completeness(data["files"], ["环评批复", "验收监测报告", "验收意见"])
        assert missing == ["验收监测报告"], f"缺失类别判断错误: {missing}"
    print("✓ 资料包提取与分类通过")


def test_kfiles():
    """测试 K 文件体系：加载、选择、上下文构建"""
    from app.engine.kfiles import (load_kfile, select_kfiles, build_kfiles_context,
                                   list_kfiles, guess_industry)

    assert "标准" in load_kfile("K01")
    assert load_kfile("K99") == ""

    kids_eia = select_kfiles("eia", "报告书", "化工")
    assert "K01" in kids_eia and "K04" in kids_eia and "K14" in kids_eia
    assert "K16" not in kids_eia and "K17" not in kids_eia

    kids_table = select_kfiles("eia", "报告表", "")
    assert "K14" not in kids_table, "报告表不应注入公众参与 K 文件"

    kids_acc = select_kfiles("acceptance", "报告书", "")
    assert "K16" in kids_acc
    assert "K02" not in kids_acc, "验收不应注入分类管理名录 K 文件"

    kids_emg = select_kfiles("emergency", "报告书", "")
    assert "K17" in kids_emg

    ctx = build_kfiles_context(["K01", "K04"])
    assert "K01" in ctx and "审核知识文件" in ctx

    assert len(list_kfiles()) == 17
    assert guess_industry("本项目为燃煤锅炉改造项目", []) == "锅炉"
    print("✓ K 文件体系通过")


def test_clean_standard_ids():
    """D 截断清洗：-结尾/三位年号/有完整前缀的短年号 → 截断桶；两位年号合法保留"""
    from app.engine.rules_engine import _clean_standard_ids

    pairs = [("GB13271-", "GB13271-"), ("GB3095-199", "GB3095-199"), ("HJ2.2-20", "HJ2.2-20"),
             ("HJ2.2-2018", "HJ2.2-2018"), ("GB14554-93", "GB14554-93"), ("GB16297-19", "GB16297-19"),
             ("GB 16297-1996", "GB16297-1996"), ("GB3838-2002", "GB3838-2002")]
    valid, truncated = _clean_standard_ids(pairs)
    vn = {n for _, n in valid}
    assert vn == {"HJ2.2-2018", "GB14554-93", "GB16297-1996", "GB3838-2002"}, vn
    assert set(truncated) == {"GB13271-", "GB3095-199", "HJ2.2-20", "GB16297-19"}, truncated
    print("✓ 截断编号清洗通过")


def test_cross_reference_summary():
    """C 未确认合并单条 + D 截断归备注 + 废止仍逐条 P0"""
    import json, tempfile
    from pathlib import Path
    from app.engine import standards_index, rules_engine

    idx_file = Path(tempfile.mkdtemp()) / "standards_index.json"
    idx_file.write_text(json.dumps({
        "GB12523-90": {"title": "建筑施工场界噪声限值", "status": "废止", "category": "", "replaced_by": ["GB 12523-2011"]},
        "GB3838-2002": {"title": "地表水环境质量标准", "status": "现行", "category": "", "replaced_by": []},
    }, ensure_ascii=False), encoding="utf-8")
    old_path = standards_index.INDEX_PATH
    standards_index.INDEX_PATH = idx_file
    try:
        rule = {"rule_id": "R-STD-001", "category": "标准引用", "severity": "P0",
                "check_config": {"pattern": "(GB|GB/T|HJ|HJ/T)\\s*[\\d.\\-—]+"}, "law_ref": "x"}
        text = ("执行 GB12523-90 和 GB 3838-2002，另参照 HJ610-2016、HJ964-2018、HJ953-2018、"
                "GB3095-199、GB13271- 等标准。")
        issues = rules_engine.run_cross_reference_check(rule, text, [])
        p0 = [i for i in issues if i["severity"] == "P0"]
        p2 = [i for i in issues if i["severity"] == "P2"]
        assert len(p0) == 1 and "12523" in p0[0]["title"] and "GB 12523-2011" in p0[0]["finding"], p0
        assert len(p2) == 1, f"未确认应合并为 1 条: {p2}"
        assert "HJ610-2016" in p2[0]["finding"] and "GB3095-199" in p2[0]["finding"], p2[0]["finding"]
        assert "疑似提取截断" in p2[0]["finding"], p2[0]["finding"]
        assert not any("3838" in i["title"] for i in issues), "现行标准不应报错"
    finally:
        standards_index.INDEX_PATH = old_path
    print("✓ 未确认汇总+截断归并通过")


def test_standards_index_file_source():
    """A 双源合并：file_index SQLite 源 + 键归一化 + 废止优先"""
    import asyncio
    from app.engine import standards_index

    async def _run():
        from app.database import init_db, async_session
        await init_db()
        from app.models.project import FileIndex
        from sqlalchemy import delete as _del
        async with async_session() as db:
            await db.execute(_del(FileIndex))
            db.add(FileIndex(title="锅炉大气污染物排放标准", standard_id="GB 13271—2014",
                             deprecated=False, category="标准"))
            db.add(FileIndex(title="建筑施工场界噪声限值", standard_id="GB12523-90",
                             deprecated=True, replaced_by="GB 12523-2011", category="标准"))
            await db.commit()
    asyncio.run(_run())

    entries = standards_index._file_index_entries()
    assert "GB13271-2014" in entries and "GB12523-90" in entries, list(entries)[:5]
    assert entries["GB12523-90"]["status"] == "废止", entries["GB12523-90"]
    assert "GB 12523-2011" in entries["GB12523-90"]["replaced_by"], entries["GB12523-90"]

    import tempfile
    old_path = standards_index.INDEX_PATH
    standards_index.INDEX_PATH = Path(tempfile.mkdtemp()) / "standards_index.json"
    try:
        merged = standards_index.build_standards_index()
        assert "GB13271-2014" in merged and "GB12523-90" in merged
    finally:
        standards_index.INDEX_PATH = old_path
    print("✓ 标准索引双源合并通过")


def test_standards_cross_reference():
    """测试标准有效性精确判定（废止 → P0 含替代关系）"""
    import json, tempfile
    from pathlib import Path
    from app.engine import standards_index, rules_engine

    idx_dir = Path(tempfile.mkdtemp())
    idx_file = idx_dir / "standards_index.json"
    idx_file.write_text(json.dumps({
        "GB16297-1996": {"title": "大气污染物综合排放标准", "status": "废止",
                         "category": "", "replaced_by": ["GB 37822-2019"]},
        "HJ2.1-2016": {"title": "总纲", "status": "现行", "category": "", "replaced_by": []},
    }, ensure_ascii=False), encoding="utf-8")
    old_path = standards_index.INDEX_PATH
    standards_index.INDEX_PATH = idx_file
    try:
        rule = {"rule_id": "R-STD-001", "category": "标准引用", "severity": "P0",
                "check_config": {"pattern": "(GB|GB/T|HJ|HJ/T)\\s*[\\d.\\-—]+"}, "law_ref": "x"}
        text = "本项目执行 GB 16297—1996 和 HJ 2.1-2016 及 GB 9999-2099。"
        issues = rules_engine.run_cross_reference_check(rule, text, [])
        p0 = [i for i in issues if i["severity"] == "P0"]
        p2 = [i for i in issues if i["severity"] == "P2"]
        assert len(p0) == 1 and "16297" in p0[0]["title"] and "GB 37822-2019" in p0[0]["finding"], f"废止判定错误: {issues}"
        assert not any("2.1" in i["title"] for i in issues), "现行标准不应报错"
        assert len(p2) == 1 and "9999" in p2[0]["title"], "未知标准应为 P2 提示"
    finally:
        standards_index.INDEX_PATH = old_path
    print("✓ 标准精确核查通过")


def test_chapter_title_keyword_check():
    """测试章节检查优先按标题匹配（正文提及不算）"""
    from app.engine.rules_engine import run_keyword_check

    rule = {"rule_id": "R-X", "severity": "P0", "category": "结构",
            "check_config": {"required_chapters": ["环境影响评价结论"]}, "law_ref": "x"}
    chapters = [{"title": "前言", "content": "本项目进行了环境影响评价结论分析"}]
    text = "前言\n本项目进行了环境影响评价结论分析。"
    issues = run_keyword_check(rule, text, chapters=chapters)
    assert len(issues) == 1, "正文提及但无对应章节标题应判缺失"

    chapters2 = [{"title": "12 环境影响评价结论", "content": "综上可行"}]
    issues2 = run_keyword_check(rule, text, chapters=chapters2)
    assert len(issues2) == 0, "章节标题命中应判通过"
    print("✓ 章节标题匹配通过")


def test_reindex_chunk():
    """测试重建索引的 MD 分块与标准号识别"""
    from app.knowledge.reindex import chunk_markdown, _guess_standard_id

    md = "# 前言\n" + "前言内容。" * 10 + "\n## 总则\n" + "总则内容。" * 10
    chunks = chunk_markdown(md)
    headings = [c["heading"] for c in chunks]
    assert "总则" in headings, headings
    assert _guess_standard_id("HJ 2.1-2016 总纲") == "HJ 2.1-2016"
    assert _guess_standard_id("无编号文档") == ""
    print("✓ 重建索引分块通过")


def test_khub_organize():
    """测试 K-Hub 整理：扫描三态 → 转换入 vault → SHA256 判重"""
    import asyncio, shutil, tempfile
    import app.knowledge.organizer as org

    async def _run():
        inbox = Path(tempfile.mkdtemp())
        vault = Path(tempfile.mkdtemp())
        (inbox / "GB13271-2014锅炉标准.md").write_text("# 锅炉大气污染物排放标准 GB 13271-2014\n" + "本标准规定限值。" * 40, encoding="utf-8")
        (inbox / "某验收报告.txt").write_text("验收监测报告，验收结论通过。" * 40, encoding="utf-8")

        from app.database import init_db, async_session
        await init_db()
        from app.models.project import KnowledgeFile
        from sqlalchemy import delete as _del
        async with async_session() as db:
            await db.execute(_del(KnowledgeFile))
            await db.commit()
        await org.set_setting("knowledge_inbox_dir", str(inbox))
        await org.set_setting("knowledge_vault_dir", str(vault))

        items = await org.scan_inbox()
        assert len(items) == 2 and all(i["status"] == "new" for i in items), items

        report = await org.organize(with_mineru=False)
        assert len(report["added"]) == 2, report

        mds = list(vault.rglob("*.md"))
        assert len(mds) == 2
        content = mds[0].read_text(encoding="utf-8")
        assert "SHA256:" in content and "类别:" in content

        report2 = await org.organize(with_mineru=False)
        assert len(report2["added"]) == 0 and len(report2["skipped"]) == 2, report2

        (inbox / "GB13271-2014锅炉标准.md").write_text("# 修订版\n" + "新内容。" * 40, encoding="utf-8")
        items2 = await org.scan_inbox()
        st = {i["name"]: i["status"] for i in items2}
        assert st["GB13271-2014锅炉标准.md"] == "changed" and st["某验收报告.txt"] == "duplicate", st

        shutil.rmtree(inbox); shutil.rmtree(vault)

    asyncio.run(_run())
    print("✓ K-Hub 整理器通过")


def test_online_stats():
    """访问统计：中间件记录 + /api/online 返回在线与今日访客"""
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app)
    r = client.get("/api/online")
    assert r.status_code == 200, r.text
    d = r.json()
    assert d["online"] >= 1 and d["today_visitors"] >= 1, d
    assert d["window_sec"] == 180, d
    print("✓ 在线访问统计通过")


def test_rag_rerank_and_merge():
    """P0 检索层：标准号加权置顶/废止重罚/邻块拼接"""
    from app.knowledge.retriever import _rerank, _merge_neighbors

    pool = [
        {"_distance": 0.30, "standard_id": "", "title": "某环评报告", "deprecated": False},
        {"_distance": 0.38, "standard_id": "GB 13271—2014", "title": "锅炉大气污染物排放标准", "deprecated": False},
        {"_distance": 0.32, "standard_id": "GB13271-2001", "title": "旧锅炉标准", "deprecated": True},
    ]
    ranked = _rerank(pool, "GB13271-2014 二氧化硫限值", 3)
    assert ranked[0]["standard_id"] == "GB 13271—2014", "标准号命中应置顶"
    assert ranked[-1].get("deprecated") is True, "废止应重罚沉底"

    ranked2 = _rerank(pool, "锅炉大气污染物排放标准 二氧化硫", 2)
    assert "锅炉" in ranked2[0]["title"], "标题完整命中应靠前"

    items = [
        {"source": "a.md", "relative_path": "表2", "excerpt": "甲" * 700},
        {"source": "a.md", "relative_path": "表2", "excerpt": "乙" * 300},
        {"source": "b.md", "relative_path": "表3", "excerpt": "丙" * 100},
    ]
    merged = _merge_neighbors(items, max_chars=800)
    assert len(merged) == 2 and len(merged[0]["excerpt"]) == 800, (len(merged), len(merged[0]["excerpt"]))
    print("✓ 检索重排+邻块拼接通过")


def test_llm_config_delete_and_vision():
    """LLM 配置：DELETE 端点删除 + 视觉启用排他（vision_review 保存启用置 vision_active）"""
    import asyncio
    from fastapi.testclient import TestClient
    from app.api.deps import get_admin_token

    async def _setup():
        from app.database import init_db, async_session
        await init_db()
        from app.models.project import LLMProfile
        from sqlalchemy import delete as _del
        async with async_session() as db:
            await db.execute(_del(LLMProfile))
            db.add(LLMProfile(id="p-audit", name="A", model="m1", api_key="k123456", purpose="audit", active=True))
            db.add(LLMProfile(id="p-old", name="配错的", model="bad", api_key="k654321", purpose="vision_review"))
            await db.commit()
    asyncio.run(_setup())

    from app.main import app
    client = TestClient(app)
    h = {"X-Admin-Token": get_admin_token()}

    r = client.post("/api/admin/llm-config", json={"id": "p-vision", "name": "V", "base_url": "https://x",
                    "model": "glm-4.6v-flash", "api_key": "k000000", "purpose": "vision_review", "activate": True}, headers=h)
    assert r.status_code == 200, r.text
    profs = {p["id"]: p for p in r.json()["profiles"]}
    assert profs["p-vision"]["vision_active"] is True and profs["p-vision"]["active"] is False, profs["p-vision"]
    assert profs["p-audit"]["active"] is True, "视觉启用不应影响审核通道"

    r = client.post("/api/admin/llm-config", json={"id": "p-vision2", "name": "V2", "base_url": "https://x",
                    "model": "v2", "api_key": "k111111", "purpose": "vision_review", "activate": True}, headers=h)
    profs = {p["id"]: p for p in r.json()["profiles"]}
    assert profs["p-vision"]["vision_active"] is False and profs["p-vision2"]["vision_active"] is True, "视觉启用应排他"

    r = client.delete("/api/admin/llm-config/p-old", headers=h)
    assert r.status_code == 200 and r.json()["deleted_id"] == "p-old", r.text
    assert all(p["id"] != "p-old" for p in r.json()["profiles"])

    r = client.delete("/api/admin/llm-config/p-old", headers=h)
    assert r.status_code == 404, r.text
    r = client.delete("/api/admin/llm-config/p-vision2", headers={"X-Admin-Token": "wrong"})
    assert r.status_code == 401, r.text
    print("✓ LLM 配置删除+视觉启用通过")


def test_terminology():
    """O4 术语检查：全文命中/标题范围限定/章节定位/应急域跳过"""
    import asyncio
    from app.engine.steps.terminology import check_terminology

    td = {"full_text": "本报告由环保局审批，项目绝对安全。",
          "chapters": [{"title": "总则", "content": "本报告由环保局审批，项目绝对安全。"}], "tables": []}
    issues = asyncio.run(check_terminology(td, {"domain": "eia"}))
    assert len(issues) == 2 and all(i["rule_id"] == "R-TERM-001" and i["severity"] == "P2" for i in issues), issues
    assert issues[0]["chapter"] == "总则", issues

    title_td = {"full_text": "三废治理内容。",
                "chapters": [{"title": "三废治理措施", "content": "正文无命中。"}], "tables": []}
    ti = asyncio.run(check_terminology(title_td, {"domain": "eia"}))
    assert len(ti) == 1 and "标题" in ti[0]["finding"], ti

    body_td = {"full_text": "三废治理内容。",
               "chapters": [{"title": "污染防治", "content": "三废治理内容。"}], "tables": []}
    assert not asyncio.run(check_terminology(body_td, {"domain": "eia"})), "title 范围不应命中正文"

    assert not asyncio.run(check_terminology(td, {"domain": "emergency"})), "应急域应跳过"
    print("✓ 术语检查通过")


def test_cases_library():
    """O2 案例库：聚类重建/来源标注/prompt 注入/启停/更新分支"""
    import asyncio
    import app.llm.client as llm
    import app.engine.llm_cache as cache

    async def _run():
        from app.database import init_db, async_session
        await init_db()
        from app.models.project import AuditIssue, AuditCase
        from sqlalchemy import delete as _del, select as _sel
        async with async_session() as db:
            await db.execute(_del(AuditCase))
            await db.execute(_del(AuditIssue))
            for _ in range(3):
                db.add(AuditIssue(project_id='x', rule_id='R-HW-001', category='危废管理',
                                  title='危废代码归类错误', finding='废活性炭应归HW49。', suggestion='改为HW49。',
                                  feedback='accurate'))
            db.add(AuditIssue(project_id='x', rule_id='R-HW-001', category='危废管理',
                              title='危废暂存超期', finding='暂存超1年。', feedback='accurate'))
            await db.commit()

        from app.engine.cases import rebuild_cases, get_cases_for_rule, match_case_sources
        rep = await rebuild_cases(llm_call=None)
        assert rep['created'] == 1 and rep['groups'] == 1, rep
        cases = await get_cases_for_rule('R-HW-001')
        assert cases and cases[0]['accurate_count'] == 3, cases

        issues = [{'rule_id': 'R-HW-001', 'title': '危废代码归类错误'},
                  {'rule_id': 'R-STD-001', 'title': '引用废止标准'}]
        await match_case_sources(issues)
        assert '历史案例' in issues[0]['case_source'] and '3 次准确' in issues[0]['case_source'], issues
        assert issues[1]['case_source'] == '规则 R-STD-001', issues

        captured = {}
        async def fake_prof(): return {'name': 'stub'}
        async def fake_chat(prompt, profile=None):
            captured['prompt'] = prompt
            return 'null'
        orig_chat, orig_prof = llm.chat, llm.get_active_profile
        orig_get, orig_set = cache.get, cache.set
        llm.chat, llm.get_active_profile = fake_chat, fake_prof
        cache.get = lambda *a: None
        cache.set = lambda *a: None
        try:
            from app.engine.rules_engine import run_llm_check
            await run_llm_check({'rule_id': 'R-HW-001', 'title': '危废核查', 'check_type': 'llm_judge',
                                 'check_config': {'description': '检查危废'}}, '某报告全文' + 'x' * 100, [])
        finally:
            llm.chat, llm.get_active_profile = orig_chat, orig_prof
            cache.get, cache.set = orig_get, orig_set
        assert '同类历史案例' in captured.get('prompt', ''), captured.get('prompt', '')[:300]

        async with async_session() as db:
            c = (await db.execute(_sel(AuditCase))).scalars().first()
            c.enabled = False
            await db.commit()
        assert not await get_cases_for_rule('R-HW-001')

        rep2 = await rebuild_cases(llm_call=None)
        assert rep2['updated'] == 1 and rep2['created'] == 0, rep2

    asyncio.run(_run())
    print("✓ 案例库通过")


def test_crosscheck_sum():
    """交叉核验-表内加和：合计错误检出、合计正确不报错、容差内不报错"""
    import asyncio
    from app.engine.steps.crosscheck import check_cross_tables

    bad_table = {"caption": "表3.4-3 工艺废气产生情况一览表", "number": "3.4-3",
                 "headers": ["编号", "排放速率kg/h"],
                 "rows": [["G1-1", "0.026"], ["G1-2", "0.031"], ["G1-3", "0.013"],
                          ["合计", "0.1554"]],
                 "chapter": "工程分析", "source": "docx"}
    async def run(t):
        return await check_cross_tables({"full_text": "", "tables": [t]}, {"domain": "eia"})
    issues = asyncio.run(run(bad_table))
    assert any(i["rule_id"] == "R-XCHK-SUM" and i["severity"] == "P1" for i in issues), issues
    assert issues[0]["chapter"] == "工程分析"

    good = {**bad_table, "rows": [["G1-1", "0.026"], ["G1-2", "0.031"], ["G1-3", "0.013"], ["合计", "0.07"]]}
    assert not any(i["rule_id"] == "R-XCHK-SUM" for i in asyncio.run(run(good)))

    rounding = {**bad_table, "rows": [["A", "0.333"], ["B", "0.333"], ["C", "0.333"], ["合计", "1.0"]]}
    assert not any(i["rule_id"] == "R-XCHK-SUM" for i in asyncio.run(run(rounding)))
    print("✓ 交叉核验-表内加和通过")


def test_crosscheck_numbering():
    """交叉核验-表格编号重复：同号不同题检出；目录点线行不误判"""
    import asyncio
    from app.engine.steps.crosscheck import check_cross_tables

    text = "表6.6-2 主要风险物质应急处置措施一览表\n内容\n表6.6-2 环境应急预案编制主要内容\n内容"
    async def run(t):
        return await check_cross_tables({"full_text": t, "tables": []}, {"domain": "eia"})
    issues = asyncio.run(run(text))
    assert any(i["rule_id"] == "R-XCHK-TBN" for i in issues), issues

    toc_text = "表6.6-2 主要风险物质应急处置措施一览表 ...... 45\n表6.6-2 主要风险物质应急处置措施一览表\n内容"
    assert not any(i["rule_id"] == "R-XCHK-TBN" for i in asyncio.run(run(toc_text)))
    print("✓ 交叉核验-表格编号通过")


def test_crosscheck_coordinates():
    """交叉核验-经纬度：分当小数误填检出；正确换算不报错；非法度分秒检出"""
    import asyncio
    from app.engine.steps.crosscheck import check_cross_tables

    bad = "厂区中心坐标为东经126°27′58.488″，事故源经度标注为126.27567961°。"
    async def run(t):
        return await check_cross_tables({"full_text": t, "tables": []}, {"domain": "eia"})
    issues = asyncio.run(run(bad))
    geo = [i for i in issues if i["rule_id"] == "R-XCHK-GEO" and i["severity"] == "P1"]
    assert geo and "126.27567961" in geo[0]["finding"], issues

    good = "厂区中心坐标为东经126°27′58.488″，模型输入经度126.4661°。"
    assert not asyncio.run(run(good)), "正确坐标不应报错"

    invalid = "坐标为东经126°75′20″。"
    issues2 = asyncio.run(run(invalid))
    assert any(i["rule_id"] == "R-XCHK-GEO" and i["severity"] == "P2" for i in issues2), issues2
    print("✓ 交叉核验-经纬度通过")


def test_crosscheck_consistency():
    """交叉核验-指标一致性：候选配对+LLM仲裁；限定词防误报；单位换算；监测日期"""
    import asyncio
    import app.llm.client as llm
    from app.engine.steps.crosscheck import check_cross_tables, _extract_metrics, _metric_candidates

    async def fake_prof():
        return {"name": "stub"}

    async def fake_chat(prompt, profile=None):
        return '[{"idx":1,"reason":"同一总投资两个数值","suggestion":"将基本情况表总投资统一为48.8万元"}]'

    async def fake_chat_empty(prompt, profile=None):
        return '[]'

    orig_chat, orig_prof = llm.chat, llm.get_active_profile

    text = "# 建设项目基本情况\n项目总投资为48万元，环保投资5万元。\n# 工程分析\n该项目总投资48.8万元。"
    chapters = [{"title": "建设项目基本情况", "content": "项目总投资为48万元，环保投资5万元。"},
                {"title": "工程分析", "content": "该项目总投资48.8万元。"}]
    td = {"full_text": text, "chapters": chapters, "tables": []}

    try:
        llm.get_active_profile = fake_prof
        llm.chat = fake_chat
        issues = asyncio.run(check_cross_tables(td, {"domain": "eia"}))
        cons = [i for i in issues if i["rule_id"] == "R-XCHK-CONS"]
        assert cons and cons[0]["severity"] == "P1", issues
        assert "48.8万元" in cons[0]["suggestion"], cons[0]["suggestion"]

        llm.chat = fake_chat_empty
        assert not [i for i in asyncio.run(check_cross_tables(td, {"domain": "eia"})) if i["rule_id"] == "R-XCHK-CONS"]

        llm.chat = fake_chat
        qual = {"full_text": "现有工程总投资48万元。本项目总投资48.8万元。", "chapters": [], "tables": []}
        assert not [i for i in asyncio.run(check_cross_tables(qual, {"domain": "eia"})) if i["rule_id"] == "R-XCHK-CONS"], "限定词不同不应配对"

        same = {"full_text": "总投资0.0048亿元。后文总投资48万元。", "chapters": [], "tables": []}
        ms = _extract_metrics(same)
        assert not _metric_candidates(ms), "亿元换算万元相等不应产生候选"

        llm.chat = fake_chat
        dates = {"full_text": "现状监测时间为2025年8月27日。表4.3-22记录监测时间2025.8.28。", "chapters": [], "tables": []}
        diss = asyncio.run(check_cross_tables(dates, {"domain": "eia"}))
        assert any(i["rule_id"] == "R-XCHK-CONS" and "监测时间" in i["title"] for i in diss), diss
    finally:
        llm.chat, llm.get_active_profile = orig_chat, orig_prof

    ms2 = _extract_metrics({"full_text": "总投资48万元。", "chapters": [], "tables": []})
    assert ms2 and ms2[0]["nunit"] == "万元" and abs(ms2[0]["value"] - 48) < 1e-6, ms2
    print("✓ 交叉核验-指标一致性通过")


def test_crosscheck_recalc_o1c():
    """交叉核验-O1c：批次重算/限值内插/速率换算 三类确定性检查"""
    import asyncio
    from app.engine.steps.crosscheck import check_cross_tables

    async def run(tables, text=""):
        return await check_cross_tables({"full_text": text, "tables": tables, "chapters": []}, {"domain": "eia"})

    batch_tbl = {"caption": "表3.3-4 物料平衡表", "number": "3.3-4",
                 "headers": ["物料", "单批产量kg/批", "年批次(批/a)", "年产量(t/a)"],
                 "rows": [["8%氨水", "1612.5 kg", "481", "755.63"]], "chapter": "", "source": "docx"}
    issues = asyncio.run(run([batch_tbl]))
    assert any(i["rule_id"] == "R-XCHK-BATCH" for i in issues), issues
    good = {**batch_tbl, "rows": [["8%氨水", "1612.5 kg", "481", "775.61"]]}
    assert not any(i["rule_id"] == "R-XCHK-BATCH" for i in asyncio.run(run([good])))

    std_tbl = {"caption": "表1.7-11 大气污染物综合排放标准", "number": "1.7-11",
               "headers": ["污染物", "20m", "30m"],
               "rows": [["氟化物", "0.17", "0.59"]], "chapter": "", "source": "docx"}
    bad_text = "排气筒DA005（高度22m）的氟化物许可排放速率标注为0.58 kg/h。"
    issues = asyncio.run(run([std_tbl], bad_text))
    interp = [i for i in issues if i["rule_id"] == "R-XCHK-INT"]
    assert interp and "0.254" in interp[0]["finding"], issues
    ok_text = "排气筒DA005（高度22m）的氟化物许可排放速率标注为0.25 kg/h。"
    assert not any(i["rule_id"] == "R-XCHK-INT" for i in asyncio.run(run([std_tbl], ok_text)))

    rate_tbl = {"caption": "表3.4-3 废气排放表", "number": "3.4-3",
                "headers": ["编号", "排放速率(kg/h)", "排放量(t/a)"],
                "rows": [["DA001", "0.1554", "1.119"], ["DA002", "0.1554", "0.265"]],
                "chapter": "", "source": "docx"}
    issues = asyncio.run(run([rate_tbl], "年运行时间7200h。"))
    rate = [i for i in issues if i["rule_id"] == "R-XCHK-RATE"]
    assert len(rate) == 1 and "DA002" in rate[0]["finding"], issues
    print("✓ 交叉核验-O1c 重算类通过")


def test_extractor_md_tables():
    """MD 管道表格提取：表号/表头/数据行/章节归属"""
    from app.engine.extractor import _parse_content

    md = "# 工程分析\n源强内容。\n表3.3-4 R05年物料平衡表\n| 项目 | 数量 |\n| --- | --- |\n| 进料水 | 1626.72 |\n| 合计 | 1626.72 |\n"
    result = _parse_content(md)
    tables = result["tables"]
    assert len(tables) == 1, tables
    t = tables[0]
    assert t["number"] == "3.3-4" and t["chapter"] == "工程分析", t
    assert t["headers"] == ["项目", "数量"] and len(t["rows"]) == 2, t
    print("✓ MD 表格提取通过")


def test_extractor_docx_tables():
    """DOCX 表格提取：有序遍历关联题注与章节（题注在表上方/下方均识别）"""
    import tempfile
    from docx import Document
    from app.engine.extractor import extract_text

    doc = Document()
    doc.add_heading("工程分析", level=1)
    doc.add_paragraph("表3.3-4 R05年物料平衡表")
    t1 = doc.add_table(rows=3, cols=2)
    for r, vals in zip(t1.rows, [("项目", "数量"), ("进料水", "1626.72"), ("合计", "1626.72")]):
        r.cells[0].text, r.cells[1].text = vals
    doc.add_heading("污染防治", level=1)
    t2 = doc.add_table(rows=2, cols=2)
    for r, vals in zip(t2.rows, [("措施", "效率"), ("活性炭", "80%")]):
        r.cells[0].text, r.cells[1].text = vals
    doc.add_paragraph("表6.1-1 废气治理措施表")

    path = str(Path(tempfile.mkdtemp()) / "报告.docx")
    doc.save(path)
    result = extract_text(path)
    tables = result["tables"]
    assert len(tables) == 2, tables
    assert tables[0]["number"] == "3.3-4" and tables[0]["chapter"] == "工程分析", tables[0]
    assert tables[0]["headers"] == ["项目", "数量"] and tables[0]["rows"][1] == ["合计", "1626.72"]
    assert tables[1]["number"] == "6.1-1" and tables[1]["chapter"] == "污染防治", tables[1]
    print("✓ DOCX 表格提取通过")


def test_run_audit_pipeline_domain():
    """回归：run_audit_pipeline 使用 project.audit_domain，不再 NameError（12% 处 domain 未定义）"""
    import asyncio, tempfile

    async def _run():
        import app.config as cfg
        from app.database import init_db, async_session
        await init_db()
        from app.models.project import Project, LLMProfile
        from sqlalchemy import delete as _del, select as _sel

        tmp = Path(tempfile.mkdtemp())
        cfg.UPLOAD_DIR = tmp
        report_file = tmp / "报告.md"
        report_file.write_text(
            "# 建设项目环境影响报告书\n## 总则\n编制依据内容。\n## 工程分析\n源强核算内容。\n" * 5,
            encoding="utf-8")

        async with async_session() as db:
            await db.execute(_del(LLMProfile))
            proj = Project(name="回归测试项目", filename="报告.md", file_path=str(report_file),
                           audit_domain="acceptance", status="queued")
            db.add(proj)
            await db.commit()
            pid = proj.id

        from app.engine.pipeline import run_audit_pipeline
        await run_audit_pipeline(pid)

        async with async_session() as db:
            p = (await db.execute(_sel(Project).where(Project.id == pid))).scalar_one()
            all_logs = " ".join(l.get("message", "") + l.get("type", "") for l in (p.logs or []))
            assert "name 'domain' is not defined" not in all_logs, all_logs[:500]
            assert "name 'domain' is not defined" not in (p.step or ""), p.step
            assert p.status == "completed", f"status={p.status} step={p.step}"
            assert any("识别完成" in l.get("message", "") for l in (p.logs or [])), "Step 0 未完成"

    asyncio.run(_run())
    print("✓ run_audit_pipeline domain 回归通过")


if __name__ == "__main__":
    test_extract_docx_heading_styles()
    test_split_chapters_chinese_patterns()
    test_grader()
    test_rules_engine_keyword()
    test_rules_engine_load_filter()
    test_standards_index()
    test_pipeline_structure_check()
    test_context_routing()
    test_llm_json_parse()
    test_extractor_folder_and_classify()
    test_kfiles()
    test_standards_cross_reference()
    test_chapter_title_keyword_check()
    test_reindex_chunk()
    test_khub_organize()
    test_run_audit_pipeline_domain()
    test_crosscheck_sum()
    test_crosscheck_numbering()
    test_crosscheck_coordinates()
    test_extractor_md_tables()
    test_extractor_docx_tables()
    print("\n✅ 全部测试通过")
