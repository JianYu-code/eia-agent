"""智能报告生成器 — K文件注入 + 行业适配 + 避坑警示 + 并发生成"""
import asyncio
from datetime import datetime

from app.knowledge.retriever import search_knowledge
from app.llm.client import chat, get_active_profile
from app.engine.kfiles import build_kfiles_context

CHAPTERS_BOOK = [
    {"id": "ch01", "title": "前言", "query": "环境影响评价 前言 编制背景 技术路线",
     "kfiles": ["K07"], "warn": "前言须含任务由来、评价工作过程、主要结论三要素"},
    {"id": "ch02", "title": "总则", "query": "环境影响评价 编制依据 评价因子 评价标准 评价等级",
     "kfiles": ["K01", "K09", "K02"], "warn": "标准引用必须带年号且为现行版本；评价等级须给出判定依据"},
    {"id": "ch03", "title": "建设项目工程分析", "query": "工程分析 工艺流程 污染源源强核算 物料平衡",
     "kfiles": ["K04", "K05"], "warn": "源强必须有出处（物料衡算/系数出处/类比条件）；须含非正常工况；改扩建须三本账"},
    {"id": "ch04", "title": "环境现状调查与评价", "query": "环境现状调查 环境质量监测 评价方法",
     "kfiles": ["K10"], "warn": "监测布点/频次须满足导则；引用数据须在3年有效期内；超标须分析原因"},
    {"id": "ch05", "title": "环境影响预测与评价", "query": "环境影响预测 大气预测 地表水预测 噪声预测",
     "kfiles": ["K11"], "warn": "预测因子与评价因子一致；须叠加背景值；防护距离须计算；占位数值用【】标注"},
    {"id": "ch06", "title": "环境保护措施及其可行性论证", "query": "污染防治措施 可行性论证 技术经济分析",
     "kfiles": ["K06", "K12"], "warn": "禁止'采用成熟技术确保达标'空话；须给工艺参数与去除效率；危废须明确代码与去向"},
    {"id": "ch07", "title": "环境影响经济损益分析", "query": "环境影响经济损益 环保投资 经济效益",
     "kfiles": [], "warn": "环保投资须列出明细并与措施对应"},
    {"id": "ch08", "title": "环境管理与监测计划", "query": "环境管理 监测计划 竣工验收 排污许可",
     "kfiles": ["K13", "K10"], "warn": "监测计划须含点位/因子/频次；总量指标须说明来源"},
    {"id": "ch09", "title": "环境影响评价结论", "query": "环境影响评价结论 综合评价 建议",
     "kfiles": ["K07"], "warn": "结论须与正文预测结果一致，不得矛盾"},
]

CHAPTERS_TABLE = [
    {"id": "tb01", "title": "建设项目基本情况", "query": "报告表 基本情况 项目概况",
     "kfiles": ["K07", "K02"], "warn": "须含立项依据、建设地点、规模、投资"},
    {"id": "tb02", "title": "建设项目工程分析", "query": "工程分析 污染源 排放量",
     "kfiles": ["K04", "K05"], "warn": "须给出各污染源排放量估算及依据"},
    {"id": "tb03", "title": "区域环境质量现状、环境保护目标及评价标准", "query": "环境质量现状 保护目标 评价标准",
     "kfiles": ["K10", "K01"], "warn": "标准须带年号；敏感目标须列名称方位距离"},
    {"id": "tb04", "title": "主要环境影响和保护措施", "query": "环境影响 保护措施 达标分析",
     "kfiles": ["K06", "K12"], "warn": "措施须具体并有达标论证"},
    {"id": "tb05", "title": "环境保护措施监督检查清单", "query": "监督检查清单 验收要求",
     "kfiles": ["K16"], "warn": "清单须含措施内容、执行标准、监测要求、完成时限"},
    {"id": "tb06", "title": "结论", "query": "评价结论 污染物排放量汇总",
     "kfiles": ["K07", "K13"], "warn": "须附污染物排放量汇总表；结论与正文一致"},
]


def _chapters_for(report_type: str) -> list[dict]:
    return CHAPTERS_TABLE if report_type == "报告表" else CHAPTERS_BOOK


def _build_chapter_prompt(ch: dict, info_text: str, kb_ctx: str, kfiles_ctx: str, industry: str) -> str:
    kf_section = f"\n写作规范（K 文件，必须遵循）：\n{kfiles_ctx}\n" if kfiles_ctx else ""
    return f"""你是具有多年经验的环评工程师。请撰写环评报告"{ch['title']}"章节。

项目信息：
{info_text}

行业类别：{industry or '未明确'}
{kf_section}
参考标准（请在正文中适当引用）：
{kb_ctx}

要求：
1. 使用专业的环评术语和规范格式
2. 引用标准必须标注标准编号+年号，且只引用参考标准中出现的编号，禁止编造
3. 内容要具体，不要空洞的套话
4. 缺真实数据时用【待补：xxx】占位，禁止编造监测数值
5. 避坑警示：{ch['warn']}
6. 章节标题使用"{ch['title']}"
7. 直接输出章节正文，不要写"以下是XX章节"之类的引导语"""


async def _generate_one(ch: dict, info_text: str, industry: str, profile) -> dict:
    kb = search_knowledge(ch["query"], top_k=4)
    kb_ctx = "\n\n".join([f"[{r['title']}] {r['excerpt'][:400]}" for r in kb])
    kfiles_ctx = build_kfiles_context(ch.get("kfiles", []), max_chars=4000) if ch.get("kfiles") else ""
    prompt = _build_chapter_prompt(ch, info_text, kb_ctx, kfiles_ctx, industry)
    content = await chat(prompt, profile=profile)
    return {"id": ch["id"], "title": ch["title"], "content": content.strip()}


async def generate_report(project_info: dict, progress_callback=None) -> dict:
    """并发逐章生成。返回 {chapters: [...], report_type, industry}"""
    profile = await get_active_profile()
    if not profile:
        raise RuntimeError("未配置启用的 LLM Profile")

    report_type = project_info.get("report_type") or "报告书"
    industry = project_info.get("industry", "")
    chapters_cfg = _chapters_for(report_type)
    info_text = _format_project_info(project_info)

    results: list[dict | Exception] = []
    tasks = [asyncio.create_task(_generate_one(ch, info_text, industry, profile)) for ch in chapters_cfg]
    total = len(tasks)
    done = 0
    for fut in asyncio.as_completed(tasks):
        try:
            res = await fut
        except Exception as e:
            res = e
        results.append(res)
        done += 1
        if progress_callback:
            await progress_callback(
                int(10 + 40 * done / total),
                "智能生成中",
                f"章节生成进度 {done}/{total}",
            )

    order = {ch["id"]: i for i, ch in enumerate(chapters_cfg)}
    ok = [r for r in results if not isinstance(r, Exception)]
    failed = [r for r in results if isinstance(r, Exception)]
    ok.sort(key=lambda r: order[r["id"]])
    for r in failed:
        ok.append({"id": f"err{len(ok)}", "title": "（章节生成失败）", "content": f"[生成失败: {str(r)[:100]}]"})
    return {"chapters": ok, "report_type": report_type, "industry": industry}


def render_markdown(project_info: dict, chapters: list[dict], output_path: str):
    lines = []
    name = project_info.get("name", "建设项目环境影响报告")
    company = project_info.get("company", "")
    date_str = datetime.now().strftime("%Y年%m月")

    lines.append(f"# {name}")
    lines.append("")
    lines.append(f"**建设单位：** {company}")
    lines.append(f"**编制日期：** {date_str}")
    lines.append("")
    lines.append("---")
    lines.append("")

    for ch in chapters:
        content = (ch.get("content") or "").strip()
        title = ch.get("title", "")
        if not content or content.startswith("[生成失败"):
            lines.append(f"## {title}")
            lines.append("")
            lines.append(f"*{content or '（本章节未能生成）'}*")
            lines.append("")
            continue
        if content.startswith(f"# {title}") or content.startswith(f"## {title}"):
            lines.append(content)
        else:
            lines.append(f"## {title}")
            lines.append("")
            lines.append(content)
        lines.append("")
        lines.append("")

    md = "\n".join(lines)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(md)
    return output_path


def render_docx(project_info: dict, chapters: list[dict], output_path: str):
    """从结构化章节直接生成 DOCX"""
    from docx import Document
    doc = Document()
    doc.add_heading(project_info.get("name", "建设项目环境影响报告"), level=0)
    doc.add_paragraph(f"建设单位：{project_info.get('company', '')}")
    doc.add_paragraph(f"编制日期：{datetime.now().strftime('%Y年%m月')}")
    for ch in chapters:
        title = ch.get("title", "")
        content = (ch.get("content") or "").strip()
        doc.add_heading(title, level=1)
        for para in content.split("\n\n"):
            para = para.strip().lstrip("#").strip()
            if para:
                doc.add_paragraph(para)
    doc.add_paragraph("说明：本报告由 AI 辅助生成，其中【待补】占位数据须由编制人员核实填写后方可使用。")
    doc.save(output_path)
    return output_path


def _format_project_info(info: dict) -> str:
    return f"""项目名称：{info.get('name', '未命名')}
建设单位：{info.get('company', '')}
建设地点：{info.get('location', '')}
建设性质：{info.get('nature', '新建')}
行业类别：{info.get('industry', '')}
总投资：{info.get('investment', '')} 万元
用地面积：{info.get('area', '')} 平方米

工程概况：
{info.get('overview', '')}

主要原辅材料：
{info.get('materials', '')}

工艺流程简述：
{info.get('process', '')}

主要设备：
{info.get('equipment', '')}

废气污染源及治理措施：
{info.get('air_pollution', '')}

废水污染源及治理措施：
{info.get('water_pollution', '')}

固废产生及处置：
{info.get('solid_waste', '')}

噪声源及控制措施：
{info.get('noise', '')}

环境敏感目标：
{info.get('sensitive_targets', '')}

环境功能区划：
{info.get('env_function', '')}

补充说明：
{info.get('notes', '')}"""
