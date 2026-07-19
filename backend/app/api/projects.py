import os
import uuid
import asyncio
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks, Request
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy import select, delete
from sqlalchemy.ext.asyncio import AsyncSession
import json

from app.database import get_db, async_session
from app.config import UPLOAD_DIR, MAX_UPLOAD_BYTES
from app.models.project import Project

router = APIRouter(prefix="/api", tags=["projects"])


@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    audit_domain: str = "eia",
    db: AsyncSession = Depends(get_db),
):
    if file.size and file.size > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="文件不能超过 200MB")

    ext = os.path.splitext(file.filename or "report")[1] or ".dat"
    stored_name = f"{uuid.uuid4().hex}{ext}"
    saved_path = UPLOAD_DIR / stored_name

    content = await file.read()
    saved_path.write_bytes(content)

    project = Project(
        name=file.filename or "未命名项目",
        filename=file.filename or "",
        file_path=str(saved_path),
        file_size=len(content),
        audit_domain=audit_domain,
        logs=[{"time": datetime.now().strftime("%H:%M:%S"), "message": f"文件上传完成：{file.filename}", "type": "success"}],
    )
    db.add(project)
    await db.commit()
    await db.refresh(project)
    return {"project": project.to_dict()}


@router.post("/upload-folder")
async def upload_folder(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    form = await request.form()
    files = form.getlist("files")
    folder_name = form.get("folder_name", "folder_report")
    audit_domain = form.get("audit_domain", "eia")

    folder_dir = UPLOAD_DIR / f"{uuid.uuid4().hex}"
    folder_dir.mkdir(exist_ok=True)

    total_size = 0
    saved_files = []
    for f in files:
        if hasattr(f, "filename"):
            content = await f.read()
            rel_path = getattr(f, "filename", "unknown")
            dest = folder_dir / os.path.basename(rel_path)
            dest.write_bytes(content)
            total_size += len(content)
            saved_files.append(rel_path)

    project = Project(
        name=folder_name,
        filename=f"{folder_name}/ ({len(saved_files)}个文件)",
        file_path=str(folder_dir),
        file_size=total_size,
        audit_domain=audit_domain,
        logs=[{"time": datetime.now().strftime("%H:%M:%S"), "message": f"文件夹上传完成：{folder_name}，共 {len(saved_files)} 个文件", "type": "success"}],
    )
    db.add(project)
    await db.commit()
    await db.refresh(project)
    return {"project": project.to_dict()}


@router.get("/projects")
async def list_projects(db: AsyncSession = Depends(get_db)):
    result = await db.execute(
        select(Project).where(Project.deleted_by_user == False).order_by(Project.created_at.desc())
    )
    projects = result.scalars().all()
    return {"projects": [p.to_dict() for p in projects]}


@router.post("/projects/{project_id}/start")
async def start_audit(
    project_id: str,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    if project.status in ("running", "queued"):
        return {"message": "该项目已在审核队列中", "project": project.to_dict()}

    from app.engine.queue import enqueue
    position = await enqueue(project_id)
    await db.refresh(project)

    msg = "审核任务已启动" if position <= 1 else f"已加入队列（前面还有 {position - 1} 个任务）"
    return {"project": project.to_dict(), "message": msg, "queue_position": position}


@router.get("/projects/{project_id}/report/view")
async def view_report(project_id: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project or not project.report_path:
        raise HTTPException(status_code=404, detail="报告不存在")

    from fastapi.responses import HTMLResponse
    content = open(project.report_path, "r", encoding="utf-8").read()
    return HTMLResponse(content)


@router.get("/projects/{project_id}/report/download")
async def download_report(
    project_id: str,
    format: str = "html",
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project or not project.report_path:
        raise HTTPException(status_code=404, detail="报告不存在")

    if format == "docx":
        try:
            from docx import Document
            from app.models.project import AuditIssue

            doc = Document()
            doc.add_heading(f"审核报告 - {project.name}", level=0)
            summary = project.result_summary or {}
            if summary.get("grade"):
                doc.add_paragraph(f"质量评级：{summary['grade']}　P0：{summary.get('p0_count', 0)}　P1：{summary.get('p1_count', 0)}　P2：{summary.get('p2_count', 0)}")
            if summary.get("summary"):
                doc.add_paragraph(f"专家组总结：{summary['summary']}")
            for t in (summary.get("top3") or []):
                doc.add_paragraph(f"优先整改：{t}")

            result = await db.execute(
                select(AuditIssue).where(AuditIssue.project_id == project_id)
            )
            db_issues = result.scalars().all()

            if db_issues:
                order = {"P0": 0, "P1": 1, "P2": 2}
                db_issues = sorted(db_issues, key=lambda x: (order.get(x.severity, 3), x.rule_id))
                cur_sev = None
                for iss in db_issues:
                    if iss.severity != cur_sev:
                        cur_sev = iss.severity
                        doc.add_heading(f"{cur_sev} 问题", level=1)
                    doc.add_heading(iss.title, level=2)
                    doc.add_paragraph(f"发现：{iss.finding}")
                    if iss.evidence_location:
                        doc.add_paragraph(f"报告原文：{iss.evidence_location}")
                    if iss.reasoning:
                        doc.add_paragraph(f"AI推理：{iss.reasoning}")
                    if iss.law_ref:
                        doc.add_paragraph(f"引用法规：{iss.law_ref}")
                    if iss.suggestion:
                        doc.add_paragraph(f"建议：{iss.suggestion}")
                    meta = " ｜ ".join(filter(None, [iss.step, iss.chapter, iss.rule_id]))
                    if meta:
                        doc.add_paragraph(meta)
            else:
                from bs4 import BeautifulSoup
                html = open(project.report_path, "r", encoding="utf-8").read()
                soup = BeautifulSoup(html, "lxml")
                for el in soup.select("h1,h2,h3,p,li"):
                    tag = el.name
                    text = el.get_text(strip=True)
                    if not text:
                        continue
                    if tag == "h1":
                        doc.add_heading(text, level=1)
                    elif tag == "h2":
                        doc.add_heading(text, level=2)
                    elif tag == "h3":
                        doc.add_heading(text, level=3)
                    else:
                        doc.add_paragraph(text)

            doc.add_paragraph("免责声明：本审核报告由 AI 自动生成，为专家复核级智能审核意见，仅供参考。")
            import tempfile
            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            doc.save(tmp.name)
            tmp.close()
            return FileResponse(tmp.name, filename=f"审核报告_{project.name}.docx",
                               media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except ImportError:
            raise HTTPException(status_code=400, detail="需要安装 python-docx 和 beautifulsoup4")

    return FileResponse(project.report_path, filename=f"审核报告_{project.name}.html")


@router.delete("/projects/{project_id}")
async def delete_project(project_id: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    if project.status in ("running", "queued"):
        raise HTTPException(status_code=400, detail="审核中的项目无法删除，请先暂停")

    project.deleted_by_user = True
    await db.commit()
    return {"message": "已删除"}


@router.post("/stop")
async def stop_audit(body: dict, db: AsyncSession = Depends(get_db)):
    project_id = body.get("project_id")
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    from app.engine.queue import cancel
    action = await cancel(project_id)
    await db.refresh(project)
    msg = {"stopping": "将在当前步骤完成后停止", "dequeued": "已取消排队", "stopped": "已暂停"}.get(action, "已暂停")
    return {"project": project.to_dict(), "message": msg}


@router.get("/projects/{project_id}/stream")
async def stream_audit(project_id: str):
    async def event_generator():
        last_progress = -1
        while True:
            async with async_session() as db:
                result = await db.execute(select(Project).where(Project.id == project_id))
                project = result.scalar_one_or_none()
                if not project:
                    break
                current = json.dumps({
                    "status": project.status,
                    "progress": project.progress,
                    "step": project.step,
                    "logs": project.logs or [],
                    "issues": project.to_dict().get("issues", {}),
                })
                if project.status in ("completed", "failed", "stopped"):
                    yield f"data: {current}\n\n"
                    break
                if project.progress != last_progress:
                    yield f"data: {current}\n\n"
                    last_progress = project.progress
            await asyncio.sleep(1)

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@router.get("/projects/{project_id}/extracted-text")
async def get_extracted_text(project_id: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    from app.engine.extractor import extract_text
    text_data = extract_text(project.file_path)
    full = text_data.get("full_text", "")
    chapters = text_data.get("chapters", [])
    return {
        "length": len(full),
        "chapter_count": len(chapters),
        "chapter_titles": [c["title"] for c in chapters[:20]],
        "preview": full[:3000],
    }
